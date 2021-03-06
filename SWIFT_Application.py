# Importing needed libraries
import streamlit as st
from docx import Document
import pandas as pd
import gspread
import os
from io import BytesIO
from oauth2client.service_account import ServiceAccountCredentials
import json
from dotenv import load_dotenv
load_dotenv()
import base64

# To download as txt
def get_binary_file_downloader_html(bin_file, file_label='File'):
    with open(bin_file, 'rb') as f:
        data = f.read()
    bin_str = base64.b64encode(data).decode()
    href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
    return href

# Generic Configurations
st.set_page_config(layout="wide")
st.title("SWIFT Analysis")  # GUI header
st.markdown('''
Allows user to choose the right mitigation strategies based on provided construction activity
''')
mitigation_list = [ ]

#Pre config google
json_creds = os.getenv("GOOGLE_SHEETS_CREDS_JSON")
creds_dict = json.loads(json_creds)
creds_dict["private_key"] = creds_dict["private_key"].replace("\\\\n", "\n")
creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict)
client = gspread.authorize(creds)

@st.cache
def data_from_googlesheets():
	# Importing the data from google sheets
	# gc = gspread.service_account(filename = CRED)
	gc = gspread.authorize(creds)
	sh = gc.open_by_key("1kP9veqfsTKnhpeO5CL9A8OLFZFlT4aaYiivBQwAihfY")
	worksheet = sh.sheet1
	data = worksheet.get_all_values()
	headers = data.pop(0)
	google_Data = pd.DataFrame(data , columns=headers)
	return google_Data


# Function to save the word document
def get_docx_download_link(docx, filename):
	"""Generates a link allowing the docx file to be downloaded
	in:  document object, filename
	out: href string
	"""
	output = BytesIO()
	docx.save(output)
	docx_str = base64.b64encode(output.getvalue()).decode()
	href = f'<a href="data:application/octet-stream;base64,{docx_str}" download="{filename}">Download the report: {filename}</a>'
	return href


# Uploading the data from google sheets
db = data_from_googlesheets()
county_df = db[ [ 'County' , 'Species' , 'Question' ] ].dropna()
construction_df = db[ [ 'Construction' , 'Possible_Construction_Activity' ] ].dropna()
mitigation_df = db[ [ 'Mitigation_Species' ,
                      'Mitigation_Construction' , 'Mitigation_Id' , 'Mitigation_Description' ] ].dropna()

# Main sheet components

# Project Details
project_details_expander = st.beta_expander("Project Location Details")
with project_details_expander:
	user_col1 , user_col2 , user_col3 = st.beta_columns(3)
	with user_col1:
		cdot_input = st.text_input("CDOT Contact")
		date = st.date_input("Date")
	with user_col2:
		project_name = st.text_input("Project Name")
		location = st.text_input("Project Location")
	with user_col3:
		project_number = st.text_input("Project Number")
		sub_account_number = st.text_input("Sub Account Number")
	project_description = st.text_input("Project Description")

# Separator
st.markdown("""
			---
			# SWIFT workflow
			---
			"""
            )

# Container for filters
filter_container = st.beta_container()
with filter_container:
	filter_container_col1 , filter_container_col2 = st.beta_columns(2)
	with filter_container_col1:
		# filter_container_col1.header("Choose Filters")
		selected_county_list = st.multiselect("Select the County" , sorted(county_df.County.unique()))
		selected_species_list = st.multiselect("Select the Species" ,
		                                       sorted(county_df.query(
				                                       'County in @selected_county_list').Species.unique()))
	with filter_container_col2:
		st.markdown("## Potential Impacts")
		selected_impacts_list_value = [ st.checkbox(i , value=i) for i in county_df.query(
				'County in @selected_county_list and Species '
				'in @selected_species_list').Question.unique() ]
		selected_impacts_list = county_df.query(
				'County in @selected_county_list and Species '
				'in @selected_species_list').Question.unique()
		final_selected_impact_list_inital = selected_impacts_list_value * selected_impacts_list
		final_selected_impact_list = filter(None , final_selected_impact_list_inital)
# st.write(final_selected_impact_list)

# selected_impacts_list = st.multiselect("Select possible environment impacts", sorted(county_df.query(
# 		                                       'County in @selected_county_list and Species '
# 		                                       'in @selected_species_list').Question.unique()))

# Container for construction activity
construction_container = st.beta_container()
with construction_container:
	construction_container_col1 , construction_container_col2 = st.beta_columns((2 , 1))
	with construction_container_col1:
		cons_activity_list = st.multiselect("Select the Construction Activity" ,
		                                    sorted(construction_df.Construction.unique()))
	with construction_container_col2:
		st.markdown("## Possible Construction Activity")
		possible_construction_container = st.beta_container()
		with possible_construction_container:
			possible_activity_list = sorted(construction_df.query('Construction in '
			                                                      '@cons_activity_list').Possible_Construction_Activity.unique())
			st.text('\n'.join(possible_activity_list))

# Container for Mitigation
st.markdown('''---''')
st.markdown('# Mitigation Strategies')
if cons_activity_list:
	mitigation_df = db[ [ 'Mitigation_Species' ,
	                      'Mitigation_Construction' , 'Mitigation_Id' , 'Mitigation_Description' ] ].dropna()
	st.write("Possible Mitigation Strategies")
	custom_mitigation_df = mitigation_df.query(
			'Mitigation_Species in @selected_species_list and Mitigation_Construction in @cons_activity_list')
	mitigation_list = sorted(custom_mitigation_df.Mitigation_Description.unique())
	st.write(mitigation_list)
else:
	pass
st.markdown('''---''')

# Writing data to microsoft word
document = Document()
document.add_heading('SWIFT ANALYSIS' , 0)
document.add_paragraph(f'SWIFT analysis conducted on {date}')

# Section 1
document.add_heading(f'Project Details' , level=1)
p1 = document.add_paragraph(f'')
p1.add_run().add_break()
p1.add_run(f'CDOT Contact: {cdot_input}')
p1.add_run().add_break()
p1.add_run(f'Project Name: {project_name}')
p1.add_run().add_break()
p1.add_run(f'Project Number: {project_number}')
p1.add_run().add_break()
p1.add_run(f'Sub Account Number: {sub_account_number}')
p1.add_run().add_break()
p1.add_run(f'Project Description: {project_description}')
p1.add_run().add_break()

# Section 2
document.add_heading(f'County' , level=2)
p2 = document.add_paragraph(f'Selected County list:')
for i in selected_county_list:
	document.add_paragraph(f'{i}' , style='List Bullet')

document.add_heading(f'Species' , level=2)
p3 = document.add_paragraph(f'Selected Species list:')
for i in selected_species_list:
	document.add_paragraph(f'{i}' , style='List Bullet')

document.add_heading(f'Potential impacts' , level=2)
p4 = document.add_paragraph(f'Selected impacts list:')
for i in final_selected_impact_list:
	document.add_paragraph(f'{i}' , style='List Bullet')

# Section 3
document.add_heading(f'Construction Activities' , level=2)
p4 = document.add_paragraph(f'Selected activities list:')
for i in cons_activity_list:
	document.add_paragraph(f'{i}' , style='List Bullet')

# Section 4
document.add_heading(f'Mitigation Strategies' , level=2)
p4 = document.add_paragraph(f'Possible strategies list:')
for i in mitigation_list:
	document.add_paragraph(f'{i}' , style='List Bullet')

st.markdown('# Save Work')
st.write(f'To save the file, Please provide the filename')
filename = st.text_input("Filename: ")
if st.button("save") and filename:
	filename = filename+".docx"
	st.markdown(get_docx_download_link(document,filename), unsafe_allow_html=True)
else:
	st.write("Provide filename to save the file, if needed")